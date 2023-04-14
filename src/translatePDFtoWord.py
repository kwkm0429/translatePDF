import PyPDF2
import docx
#import requests

# PDFファイルからテキストを抽出する関数
def extract_text_from_pdf(pdf_path):
    pdf_file = open(pdf_path, 'rb')
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ''
    for page in range(len(pdf_reader.pages)):
        page_obj = pdf_reader.pages[page]
        text += page_obj.extract_text()
    pdf_file.close()
    return text
"""
# Google翻訳APIを使用してテキストを翻訳する関数
def translate_text(text, source_language='en', target_language='ja'):
    url = 'https://translation.googleapis.com/language/translate/v2'
    headers = {'content-type': 'application/json'}
    data = {
        'q': text,
        'source': source_language,
        'target': target_language,
        'format': 'text'
    }
    api_key = '<YOUR_API_KEY>' # Google Cloud ConsoleでAPIキーを生成してください
    params = {'key': api_key}
    response = requests.post(url, headers=headers, json=data, params=params)
    translated_text = response.json()['data']['translations'][0]['translatedText']
    return translated_text
"""

# PDFファイルからテキストを抽出する
pdf_text = extract_text_from_pdf('file_pdf/hello.pdf')

# テキストをWordファイルに書き込む
doc = docx.Document()
doc.add_paragraph(pdf_text)
doc.save('file_docx/hello.docx')

"""
# テキストを翻訳する
translated_text = translate_text(pdf_text)

# 翻訳されたテキストをWordファイルに書き込む
doc = docx.Document()
doc.add_paragraph(translated_text)
doc.save('output_translated.docx')
"""