#!/usr/bin/env python
# -*- coding: utf-8 -*-
import docx
import re
import sys

args = sys.argv

# Word文書のパスを受け取り、テキストを読み込んで返す
def read_word_text(filepath):
    doc = docx.Document(filepath)
    text = ''
    for para in doc.paragraphs:
        text += para.text
    return text

# 改行自動削除、ハイフンで区切られた英単語を連結する関数
def remove_newlines_and_join_hyphenated_words(text):
    # 改行を削除する
    text = text.replace('\n', ' ')
    # ハイフンで区切られた英単語を連結する
    text = re.sub(r'([a-zA-Z]+)-\s?([a-zA-Z]+)', r'\1\2', text)
    return text

# 与えられた文字列を最大5000文字の長さで改行する
def break_lines_with_blank(text):
    max_length = 4000
    lines = []
    line = ''
    for sentence in text.split('.'):
        if len(line) + len(sentence) +1 <= max_length:
            if line:
                line += ' '
            line += sentence
        else:
            lines.append(line.strip() + '.')
            line = sentence.strip()
    if line:
        lines.append(line.strip() + '.')
    return '\n'.join(lines)

word_text = read_word_text('file_docx/'+args[1]+'.docx')
word_text = remove_newlines_and_join_hyphenated_words(word_text)
word_text = break_lines_with_blank(word_text)
word_text = word_text.replace('\n', '\n\n') #改行をWordの改行に変換する

# テキストをWordファイルに書き込む
doc = docx.Document()
"""
paragraphs = word_text.split('\n\n')
for p in paragraphs:
    doc.add_paragraph(p)
    doc.add_paragraph('') #空白行を追加する
"""
doc.add_paragraph(word_text)
doc.save('file_docx/'+args[1]+'-formatted.docx')